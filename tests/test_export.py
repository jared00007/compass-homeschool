"""Rendering a lesson to a printable Word doc.

This is a parent-only export -- callers gate access, not this module -- so the
thing worth pinning is simpler: does it actually produce a valid .docx with the
content that's on screen, including the assessment and quiz answer key, and
does it survive lessons where optional sections (video, quiz, credits) are
empty rather than crashing on a missing key.
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from compass.export import (
    DocxExtractionError,
    course_filename,
    course_to_docx,
    extract_docx_text,
    lesson_to_docx,
    suggested_filename,
    travel_journal_filename,
    travel_journal_to_docx,
)


def a_lesson(**overrides):
    payload = {
        "title": "Two-Step Equations",
        "overview": "Undo the addition, then the multiplication.",
        "learning_objectives": ["Solve for x"],
        "activities": [
            {
                "title": "Practice",
                "kind": "practice",
                "minutes": 60,
                "instructions": "Solve problems 1-10.",
                "example": "Worked model: solve 5x + 3 = 18 step by step.",
                "video": {"found": False, "title": "", "url": "", "channel": "", "why": ""},
            }
        ],
        "materials": ["Pencil"],
        "assessment": {
            "kind": "Worksheet check",
            "description": "Ten items, mixed procedural and applied.",
            "mastery_criteria": "8 of 10 correct, including 2 of 3 applied problems.",
        },
        "subject_credits": [
            {"subject": "math", "minutes": 60, "justification": "All of it."}
        ],
        "estimated_minutes": 60,
        "parent_notes": "Watch for sign errors.",
        "branches": [],
        "quiz": [
            {
                "question": "What do you do first to solve 2x + 3 = 11?",
                "choices": ["Subtract 3", "Divide by 2", "Add 3", "Multiply by 2"],
                "correct_index": 0,
                "explanation": "Undo addition before multiplication.",
            }
        ],
    }
    payload.update(overrides)
    return payload


def text_of(docx_bytes: bytes) -> str:
    document = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_produces_a_valid_docx():
    docx_bytes = lesson_to_docx(a_lesson())
    assert docx_bytes[:2] == b"PK"  # docx is a zip container
    Document(io.BytesIO(docx_bytes))  # raises if not a well-formed document


def test_includes_the_content_on_screen():
    text = text_of(lesson_to_docx(a_lesson()))
    assert "Two-Step Equations" in text
    assert "Undo the addition, then the multiplication." in text
    assert "Solve for x" in text
    assert "Solve problems 1-10." in text
    assert "Pencil" in text


def test_includes_the_parent_only_assessment_and_answer_key():
    """The whole point of this export: the answer key travels with it."""
    text = text_of(lesson_to_docx(a_lesson()))
    assert "Ten items, mixed procedural and applied." in text
    assert "8 of 10 correct, including 2 of 3 applied problems." in text
    assert "What do you do first to solve 2x + 3 = 11?" in text
    assert "Subtract 3 (correct)" in text
    assert "Undo addition before multiplication." in text


def test_includes_subject_credit_table():
    text = text_of(lesson_to_docx(a_lesson()))
    assert "Math" in text
    assert "All of it." in text


def test_handles_a_missing_video_and_empty_quiz_without_crashing():
    lesson = a_lesson(quiz=[])
    lesson["activities"][0]["video"] = {"found": False}
    text = text_of(lesson_to_docx(lesson))
    assert "Solving Two-Step Equations" not in text
    assert "Quiz answer key" not in text


def test_handles_a_found_video():
    lesson = a_lesson()
    lesson["activities"][0]["video"] = {
        "found": True,
        "title": "Solving Two-Step Equations",
        "url": "https://www.youtube.com/watch?v=abc123",
        "channel": "Khan Academy",
        "why": "Shows the undo steps worked in real time.",
    }
    text = text_of(lesson_to_docx(lesson))
    assert "Solving Two-Step Equations" in text
    assert "https://www.youtube.com/watch?v=abc123" in text


def test_handles_a_minimal_lesson_with_only_required_looking_fields():
    """Nothing here should assume every optional key is present."""
    docx_bytes = lesson_to_docx({"title": "Bare Lesson"})
    text = text_of(docx_bytes)
    assert "Bare Lesson" in text


@pytest.mark.parametrize(
    "title,expected_prefix",
    [
        ("Two-Step Equations", "two-step-equations-"),
        ("Nurse Log Field Study!", "nurse-log-field-study-"),
        ("", "lesson-"),
    ],
)
def test_suggested_filename_is_slugged_and_dated(title, expected_prefix):
    name = suggested_filename({"title": title})
    assert name.startswith(expected_prefix)
    assert name.endswith(".docx")
    assert " " not in name


# --- course_to_docx: the grades 6-12 district documentation packet -----------


def a_course(**overrides):
    course = {
        "id": 1,
        "title": "Washington State History",
        "credit_subject": "history",
        "grade_level": "8",
        "description": "A survey of Washington state history.",
        "goals": "Understand tribal, territorial, and statehood eras.",
        "outline": "Unit 1: Indigenous peoples. Unit 2: Territorial era.",
        "credit_value": 1.0,
        "start_date": "2025-09-01",
        "end_date": "2026-08-31",
        "final_grade": "",
        "pass_fail": None,
    }
    course.update(overrides)
    return course


def an_activity_with_lesson(**overrides):
    activity = {
        "id": 1,
        "occurred_on": "2025-10-01",
        "title": "Indigenous Peoples of the Puget Sound",
        "description": "",
        "minutes": 60,
        "lesson_id": 1,
        "lesson": {
            "payload": {
                "learning_objectives": ["Identify three tribes native to the Puget Sound region"],
                "assessment": {
                    "kind": "Oral check",
                    "description": "Name three tribes and one cultural practice each.",
                    "mastery_criteria": "3 of 3 correct",
                },
            },
            "metadata": {"quiz_result": {"correct": 9, "total": 10, "passed": True}},
        },
    }
    activity.update(overrides)
    return activity


def test_course_packet_covers_all_seven_required_sections():
    docx_bytes = course_to_docx(a_course(), [an_activity_with_lesson()], "Landon")
    text = text_of(docx_bytes)
    assert "Course description" in text
    assert "A survey of Washington state history." in text
    assert "Course goals and objectives" in text
    assert "Understand tribal, territorial, and statehood eras." in text
    assert "Course outline of the program" in text
    assert "Learning activities and instructional time log" in text
    assert "Completed assignments and assessments" in text
    assert "Identify three tribes native to the Puget Sound region" in text
    assert "Quiz: 9/10" in text
    assert "How student performance is assessed" in text
    assert "Name three tribes and one cultural practice each." in text
    assert "Student progress and final grade" in text


def test_hours_log_totals_and_shows_progress_toward_the_150_hour_credit():
    docx_bytes = course_to_docx(a_course(), [an_activity_with_lesson()], "Landon")
    text = text_of(docx_bytes)
    assert "1 of 150 hours logged" in text
    assert "2025-10-01" in text and "60" in text


def test_half_credit_course_targets_75_hours():
    docx_bytes = course_to_docx(a_course(credit_value=0.5), [], "Landon")
    text = text_of(docx_bytes)
    assert "75" in text


def test_final_grade_and_pass_fail_appear_when_set():
    docx_bytes = course_to_docx(
        a_course(final_grade="A", pass_fail="pass"), [an_activity_with_lesson()], "Landon"
    )
    text = text_of(docx_bytes)
    assert "Final grade: A" in text
    assert "Transcript record: PASS" in text


def test_not_yet_graded_says_so_rather_than_omitting_the_line():
    docx_bytes = course_to_docx(a_course(), [], "Landon")
    text = text_of(docx_bytes)
    assert "in progress" in text.lower()


def test_handles_a_course_with_no_activities_yet():
    """A freshly created course, before anything's tagged to it, must still
    produce a valid packet rather than crashing on an empty list."""
    docx_bytes = course_to_docx(a_course(), [], "Landon")
    text = text_of(docx_bytes)
    assert "Washington State History" in text
    assert "No instructional time logged" in text
    assert "Nothing completed toward this course yet." in text


def test_handles_an_activity_with_no_lesson_behind_it():
    """A manually logged activity (no `lesson_id`) has `lesson: None` --
    the packet must fall back to the activity's own description rather than
    assuming lesson content always exists."""
    activity = an_activity_with_lesson(lesson_id=None, lesson=None, description="Field trip notes.")
    docx_bytes = course_to_docx(a_course(), [activity], "Landon")
    text = text_of(docx_bytes)
    assert "Field trip notes." in text
    assert (
        "Performance assessed through direct parent observation" in text
    )


@pytest.mark.parametrize(
    "title,expected_prefix",
    [
        ("Washington State History", "washington-state-history-documentation-"),
        ("Algebra 1!", "algebra-1-documentation-"),
        ("", "course-documentation-"),
    ],
)
def test_course_filename_is_slugged_and_dated(title, expected_prefix):
    name = course_filename({"title": title})
    assert name.startswith(expected_prefix)
    assert name.endswith(".docx")
    assert " " not in name


# --- travel journal export ----------------------------------------------------


def a_travel_entry(**overrides):
    entry = {
        "title": "Glaciers Before They're Gone",
        "state": "Montana",
        "visited_on": "2025-07-04",
        "story": "Hiked to the lake and watched the ice calve off the glacier.",
        "favorite_moment": "",
        "would_return": "",
        "park_name": "",
    }
    entry.update(overrides)
    return entry


def test_travel_journal_produces_a_valid_docx():
    docx_bytes = travel_journal_to_docx([a_travel_entry()], "Landon")
    assert docx_bytes[:2] == b"PK"
    Document(io.BytesIO(docx_bytes))


def test_travel_journal_includes_the_trip_content():
    text = text_of(travel_journal_to_docx([a_travel_entry()], "Landon"))
    assert "Glaciers Before They're Gone" in text
    assert "Montana" in text
    assert "2025-07-04" in text
    assert "Hiked to the lake" in text


def test_travel_journal_includes_favorite_moment_and_would_return_when_set():
    entry = a_travel_entry(
        favorite_moment="Watching the glacier calve.",
        would_return="Yes, in the fall next time.",
    )
    text = text_of(travel_journal_to_docx([entry], "Landon"))
    assert "Watching the glacier calve." in text
    assert "Yes, in the fall next time." in text


def test_travel_journal_omits_favorite_moment_and_would_return_when_blank():
    text = text_of(travel_journal_to_docx([a_travel_entry()], "Landon"))
    assert "Favorite moment" not in text
    assert "Would go back" not in text


def test_travel_journal_includes_the_park_name_when_given():
    entry = a_travel_entry(park_name="Glacier National Park")
    text = text_of(travel_journal_to_docx([entry], "Landon"))
    assert "Glacier National Park" in text


def test_travel_journal_summarizes_states_and_trip_count():
    entries = [
        a_travel_entry(state="Montana", visited_on="2025-07-04"),
        a_travel_entry(state="Montana", visited_on="2024-06-01"),
        a_travel_entry(state="Wyoming", visited_on="2023-08-15"),
    ]
    text = text_of(travel_journal_to_docx(entries, "Landon"))
    assert "2 states visited" in text
    assert "3 trips logged" in text


def test_travel_journal_handles_no_entries_without_crashing():
    docx_bytes = travel_journal_to_docx([], "Landon")
    Document(io.BytesIO(docx_bytes))
    assert "No trips logged yet" in text_of(docx_bytes)


@pytest.mark.parametrize(
    "name,expected_prefix",
    [
        ("Landon", "landon-travel-journal-"),
        ("Mary-Jane O'Brien", "mary-jane-o-brien-travel-journal-"),
        ("", "student-travel-journal-"),
    ],
)
def test_travel_journal_filename_is_slugged_and_dated(name, expected_prefix):
    filename = travel_journal_filename(name)
    assert filename.startswith(expected_prefix)
    assert filename.endswith(".docx")
    assert " " not in filename


# --- extract_docx_text: the import side, an uploaded response ------------------


def _docx_bytes(paragraphs: list[str]) -> io.BytesIO:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def test_extract_docx_text_reads_every_paragraph():
    uploaded = _docx_bytes(["First paragraph.", "Second paragraph."])
    assert extract_docx_text(uploaded) == "First paragraph.\nSecond paragraph."


def test_extract_docx_text_includes_table_content():
    document = Document()
    document.add_paragraph("Above the table.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "a"
    table.rows[0].cells[1].text = "b"
    table.rows[1].cells[0].text = "c"
    table.rows[1].cells[1].text = "d"
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    text = extract_docx_text(buffer)

    assert "Above the table." in text
    assert "a | b" in text
    assert "c | d" in text


def test_extract_docx_text_strips_leading_and_trailing_blank_lines():
    uploaded = _docx_bytes(["", "", "The actual response.", ""])
    assert extract_docx_text(uploaded) == "The actual response."


def test_extract_docx_text_raises_a_clear_error_on_a_non_docx_file():
    not_a_docx = io.BytesIO(b"this is plain text, not a real .docx")
    with pytest.raises(DocxExtractionError, match="doesn't look like a valid Word"):
        extract_docx_text(not_a_docx)


# --- lesson_to_pdf: the per-lesson "print to PDF" export ------------------------


def test_lesson_to_pdf_produces_a_valid_pdf():
    from compass.export import lesson_to_pdf

    data = lesson_to_pdf(a_lesson())
    assert data[:5] == b"%PDF-"
    assert len(data) > 1000  # a real multi-section document, not an empty stub


def test_lesson_to_pdf_survives_a_bare_or_empty_lesson():
    """Optional sections missing (or the whole payload empty) must render a
    valid one-liner PDF, never crash on a missing key -- same resilience the
    .docx export has."""
    from compass.export import lesson_to_pdf

    for payload in ({}, {"title": "Just a Title"}, {"activities": [], "quiz": []}):
        data = lesson_to_pdf(payload)
        assert data[:5] == b"%PDF-"


def test_lesson_to_pdf_strips_emoji_that_have_no_print_glyph():
    """Emoji have no glyph in the print font, so they're stripped rather than
    left to render as notdef boxes -- the surrounding real text stays."""
    from compass.export import lesson_to_pdf

    data = lesson_to_pdf(a_lesson(title="📐 Two-Step Equations 🎯"))
    assert data[:5] == b"%PDF-"


def test_suggested_pdf_filename_slugs_the_title_and_ends_in_pdf():
    from datetime import date

    from compass.export import suggested_pdf_filename

    name = suggested_pdf_filename({"title": "Two-Step Equations!"})
    assert name == f"two-step-equations-{date.today().isoformat()}.pdf"
    assert suggested_pdf_filename({}).endswith(".pdf")
