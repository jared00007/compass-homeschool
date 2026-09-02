"""Turn a lesson payload -- or a whole course -- into a printable Word document,
and the other direction: pull plain text back out of a .docx he hands in.

The lesson export comes in two cuts, chosen by the caller via `lesson_to_pdf`'s
`parent` flag (the .docx export stays parent-only). The parent cut includes
everything `render_lesson`'s parent view shows, assessment and quiz answer key
included -- it exists because reading an assessment off a laptop screen while
scoring a kid's paper worksheet is awkward, a printed page isn't. The student
cut (`parent=False`) mirrors `render_lesson`'s student view exactly: the
lesson itself (overview, objectives, materials, activities) with the answer
key, assessment, parent notes, and credit left out, so Landon can print his
own copy of a lesson off his board without it carrying anything he isn't meant
to see. Nothing here re-checks who's asking -- the caller passes the right
`parent` value for the context it's rendering in.

The course export exists for a different, higher-stakes reason: Sumner-Bonney
Lake requires this exact documentation set, per course, before a grade 6-12
course counts toward the diploma. `course_to_docx` produces one packet per
course covering all seven required pieces, built from data the app already
has (the course's own record, plus every activity/lesson tagged to it) rather
than anything re-typed for the district.

`extract_docx_text` is the import side: some kids would rather write in Word
than in a browser text box. It feeds the exact same `response` string every
other part of the writing-response flow already works with (word-count
checks, the AI "check my work" pass, parent review) -- there's no separate
review path for an uploaded doc, just a different way of getting text in.
"""

from __future__ import annotations

import io
import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from compass import config, subjects


def suggested_filename(lesson: dict[str, Any]) -> str:
    """A readable .docx filename: the lesson title, slugged, plus today's date."""
    title = lesson.get("title") or "lesson"
    slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-") or "lesson"
    return f"{slug}-{date.today().isoformat()}.docx"


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def lesson_to_docx(lesson: dict[str, Any]) -> bytes:
    """Render a lesson payload to a .docx file, returned as bytes."""
    document = Document()
    style = document.styles["Normal"]
    style.font.size = Pt(11)

    document.add_heading(lesson.get("title") or "Lesson", level=1)
    if lesson.get("overview"):
        document.add_paragraph(lesson["overview"])

    objectives = lesson.get("learning_objectives") or []
    if objectives:
        document.add_heading("Learning objectives", level=2)
        _add_bullets(document, objectives)

    materials = lesson.get("materials") or []
    if materials:
        document.add_heading("Materials", level=2)
        _add_bullets(document, materials)

    activities = lesson.get("activities") or []
    if activities:
        document.add_heading("Activities", level=2)
        for index, activity in enumerate(activities, start=1):
            heading = (
                f"{index}. {activity.get('title', 'Activity')} "
                f"({activity.get('kind', '')}, {activity.get('minutes', 0)} min)"
            )
            document.add_heading(heading, level=3)
            video = activity.get("video") or {}
            if video.get("found") and video.get("url"):
                watch = document.add_paragraph()
                watch.add_run(f"Video: {video.get('title') or 'Watch'}").bold = True
                document.add_paragraph(video["url"])
                if video.get("why"):
                    document.add_paragraph(video["why"]).runs[0].italic = True
            if activity.get("example"):
                worked = document.add_paragraph()
                worked.add_run("Here's how: ").bold = True
                worked.add_run(activity["example"])
            if activity.get("instructions"):
                document.add_paragraph(activity["instructions"])

    assessment = lesson.get("assessment") or {}
    if assessment:
        document.add_heading("Assessment", level=2)
        if assessment.get("kind"):
            document.add_paragraph(assessment["kind"]).runs[0].bold = True
        if assessment.get("description"):
            document.add_paragraph(assessment["description"])
        if assessment.get("mastery_criteria"):
            criteria = document.add_paragraph()
            criteria.add_run("Mastery: ").bold = True
            criteria.add_run(assessment["mastery_criteria"])

    quiz = lesson.get("quiz") or []
    if quiz:
        document.add_heading("Quiz answer key", level=2)
        for index, item in enumerate(quiz, start=1):
            document.add_paragraph(f"{index}. {item.get('question', '')}").runs[0].bold = True
            correct_index = item.get("correct_index")
            for choice_index, choice in enumerate(item.get("choices") or []):
                marker = " (correct)" if choice_index == correct_index else ""
                document.add_paragraph(f"{choice}{marker}", style="List Bullet")
            if item.get("explanation"):
                explanation = document.add_paragraph()
                explanation.add_run(item["explanation"]).italic = True

    if lesson.get("parent_notes"):
        document.add_heading("Notes for the parent", level=2)
        document.add_paragraph(lesson["parent_notes"])

    credits = lesson.get("subject_credits") or []
    if credits:
        document.add_heading("Subject credit", level=2)
        table = document.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        header[0].text, header[1].text, header[2].text = "Subject", "Minutes", "Why"
        for credit in credits:
            row = table.add_row().cells
            row[0].text = subjects.label(credit.get("subject", ""))
            row[1].text = str(credit.get("minutes", ""))
            row[2].text = credit.get("justification", "")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_PDF_FONTS_REGISTERED = False


def _register_pdf_fonts() -> None:
    """Register the bundled DejaVu Sans (regular + bold) with reportlab, once.
    reportlab's built-in fonts are Latin-1 only -- a curly quote or accent in a
    lesson would come out as a black box. DejaVu is a full-Unicode TTF, shipped
    in compass/assets/fonts/ so the PDF export needs no system font and no
    LibreOffice, just `pip install reportlab`."""
    global _PDF_FONTS_REGISTERED
    if _PDF_FONTS_REGISTERED:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fonts_dir = Path(__file__).resolve().parent / "assets" / "fonts"
    pdfmetrics.registerFont(TTFont("DejaVu", str(fonts_dir / "DejaVuSans.ttf")))
    pdfmetrics.registerFont(TTFont("DejaVu-Bold", str(fonts_dir / "DejaVuSans-Bold.ttf")))
    pdfmetrics.registerFontFamily(
        "DejaVu", normal="DejaVu", bold="DejaVu-Bold",
        italic="DejaVu", boldItalic="DejaVu-Bold",
    )
    _PDF_FONTS_REGISTERED = True


# Color emoji have no glyphs in DejaVu (or any print font) -- left in, they draw
# as notdef boxes, so strip them from the printable copy. The lesson's actual
# text (objectives, instructions, ...) carries no emoji; they're UI chrome.
_PDF_EMOJI_RE = re.compile(
    "[\U0001F000-\U0001FAFF\U00002600-\U000027BF\U00002B00-\U00002BFF"
    "\U0001F1E6-\U0001F1FF\U0000FE00-\U0000FE0F\U0000200D\U00002190-\U000021FF]"
)


def _pdf_text(value: Any) -> str:
    """Emoji-stripped, XML-escaped text safe to drop into a reportlab
    Paragraph (whose own markup would otherwise choke on a bare & or <)."""
    from xml.sax.saxutils import escape

    return escape(_PDF_EMOJI_RE.sub("", str(value)).strip())


def _pdf_split_blocks(text: Any) -> list[str]:
    """Multi-paragraph prose -> a list of paragraph strings, structure kept: a
    blank line starts a new paragraph, a single newline becomes an XML `<br/>`
    line break, and the whole thing is XML-escaped so it's safe to drop into a
    reportlab Paragraph. Pulled out of `lesson_to_pdf` so the structure logic
    is unit-testable without building a PDF -- it exists because reportlab
    otherwise collapses every newline to a single space, flattening a
    structured assignment into one run-on blurb (reported: "it turns paragraphs
    into blurbs and lossing the structure of assignment")."""
    from xml.sax.saxutils import escape

    cleaned = _PDF_EMOJI_RE.sub("", str(text)).replace("\r\n", "\n").strip()
    if not cleaned:
        return []
    return [escape(block).replace("\n", "<br/>") for block in re.split(r"\n\s*\n", cleaned)]


def suggested_pdf_filename(lesson: dict[str, Any]) -> str:
    """A readable .pdf filename: the lesson title, slugged, plus today's date."""
    title = lesson.get("title") or "lesson"
    slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-") or "lesson"
    return f"{slug}-{date.today().isoformat()}.pdf"


def lesson_to_pdf(lesson: dict[str, Any], *, parent: bool = True) -> bytes:
    """Render a lesson payload to a print-ready PDF, returned as bytes -- the
    same structure lesson_to_docx builds (title, overview, objectives,
    materials, activities, assessment, quiz key, parent notes, credit), so a
    parent can print any one lesson for paper work or the record. reportlab is
    imported lazily so a missing install only disables this button, never the
    whole app.

    `parent` gates the same sections `render_lesson` gates: with `parent=False`
    the assessment, quiz answer key, parent notes, and subject credit are all
    left out, giving Landon a clean printable copy of a lesson (overview,
    objectives, materials, activities) with nothing he isn't meant to see."""
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    _register_pdf_fonts()

    title_style = ParagraphStyle(
        "cp_title", fontName="DejaVu-Bold", fontSize=18, leading=22, spaceAfter=8,
        alignment=TA_LEFT,
    )
    h2 = ParagraphStyle(
        "cp_h2", fontName="DejaVu-Bold", fontSize=13, leading=16,
        spaceBefore=12, spaceAfter=4,
    )
    h3 = ParagraphStyle(
        "cp_h3", fontName="DejaVu-Bold", fontSize=11.5, leading=15,
        spaceBefore=8, spaceAfter=2,
    )
    body = ParagraphStyle("cp_body", fontName="DejaVu", fontSize=11, leading=15, spaceAfter=4)

    def bullets(items: list[Any]) -> ListFlowable:
        return ListFlowable(
            [ListItem(Paragraph(_pdf_text(i), body), leftIndent=12) for i in items],
            bulletType="bullet", start="•", leftIndent=14,
        )

    def paras(text: Any, *, prefix_html: str = "") -> list[Any]:
        """Multi-paragraph prose -> a list of Paragraph flowables that keep the
        writing's structure (see _pdf_split_blocks). An optional bold
        `prefix_html` leads the first paragraph ("Here's how:", "Mastery:")."""
        out: list[Any] = []
        for i, html in enumerate(_pdf_split_blocks(text)):
            if i == 0 and prefix_html:
                html = f"{prefix_html}{html}"
            out.append(Paragraph(html, body))
        return out

    flow: list[Any] = [Paragraph(_pdf_text(lesson.get("title") or "Lesson"), title_style)]
    if lesson.get("overview"):
        flow += paras(lesson["overview"])

    if lesson.get("learning_objectives"):
        flow += [Paragraph("Learning objectives", h2), bullets(lesson["learning_objectives"])]
    if lesson.get("materials"):
        flow += [Paragraph("Materials", h2), bullets(lesson["materials"])]

    activities = lesson.get("activities") or []
    if activities:
        flow.append(Paragraph("Activities", h2))
        for index, activity in enumerate(activities, start=1):
            heading = (
                f"{index}. {activity.get('title', 'Activity')} "
                f"({activity.get('kind', '')}, {activity.get('minutes', 0)} min)"
            )
            flow.append(Paragraph(_pdf_text(heading), h3))
            video = activity.get("video") or {}
            if video.get("found") and video.get("url"):
                flow.append(Paragraph(f"<b>Video:</b> {_pdf_text(video.get('title') or 'Watch')}", body))
                flow.append(Paragraph(_pdf_text(video["url"]), body))
                if video.get("why"):
                    flow.append(Paragraph(f"<i>{_pdf_text(video['why'])}</i>", body))
            if activity.get("example"):
                flow += paras(activity["example"], prefix_html="<b>Here's how:</b> ")
            if activity.get("instructions"):
                flow += paras(activity["instructions"])

    assessment = lesson.get("assessment") or {}
    if assessment and parent:
        flow.append(Paragraph("Assessment", h2))
        if assessment.get("kind"):
            flow.append(Paragraph(f"<b>{_pdf_text(assessment['kind'])}</b>", body))
        if assessment.get("description"):
            flow += paras(assessment["description"])
        if assessment.get("mastery_criteria"):
            flow += paras(assessment["mastery_criteria"], prefix_html="<b>Mastery:</b> ")

    quiz = lesson.get("quiz") or []
    if quiz and parent:
        flow.append(Paragraph("Quiz answer key", h2))
        for index, item in enumerate(quiz, start=1):
            flow.append(Paragraph(f"<b>{index}. {_pdf_text(item.get('question', ''))}</b>", body))
            correct_index = item.get("correct_index")
            choices = item.get("choices") or []
            flow.append(
                ListFlowable(
                    [
                        ListItem(
                            Paragraph(
                                _pdf_text(choice)
                                + (" <b>(correct)</b>" if choice_index == correct_index else ""),
                                body,
                            ),
                            leftIndent=12,
                        )
                        for choice_index, choice in enumerate(choices)
                    ],
                    bulletType="bullet", start="•", leftIndent=14,
                )
            )
            if item.get("explanation"):
                flow.append(Paragraph(f"<i>{_pdf_text(item['explanation'])}</i>", body))

    if lesson.get("parent_notes") and parent:
        flow += [Paragraph("Notes for the parent", h2), *paras(lesson["parent_notes"])]

    credits = lesson.get("subject_credits") or []
    if credits and parent:
        flow.append(Paragraph("Subject credit", h2))
        for credit in credits:
            line = (
                f"<b>{_pdf_text(subjects.label(credit.get('subject', '')))}</b> — "
                f"{_pdf_text(credit.get('minutes', ''))} min"
            )
            if credit.get("justification"):
                line += f": {_pdf_text(credit['justification'])}"
            flow.append(Paragraph(line, body))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.8 * inch, bottomMargin=0.8 * inch,
        title=(lesson.get("title") or "Lesson"),
    )
    doc.build(flow + [Spacer(1, 2)])
    return buffer.getvalue()


def travel_journal_filename(student_name: str) -> str:
    """A readable .docx filename: the student's name, slugged, plus today's date."""
    slug = re.sub(r"[^a-z0-9]+", "-", student_name.lower()).strip("-") or "student"
    return f"{slug}-travel-journal-{date.today().isoformat()}.docx"


def travel_journal_to_docx(entries: list[dict[str, Any]], student_name: str) -> bytes:
    """Render the whole travel journal as one printable keepsake, newest trip
    first -- meant to be kept and added to every year, the same way the
    journal itself is meant to, not a one-time snapshot.

    Each entry may carry an optional `park_name` (already resolved by the
    caller, since park lookups live in compass.national_parks and this
    module deliberately stays independent of feature-specific lookups --
    same reasoning course_to_docx keeps subject-label lookups in
    compass.subjects rather than duplicating them here).
    """
    document = Document()
    style = document.styles["Normal"]
    style.font.size = Pt(11)

    document.add_heading(f"{student_name}'s Travels", level=1)
    if entries:
        states = {e["state"] for e in entries}
        summary = document.add_paragraph()
        summary.add_run(
            f"{len(states)} state{'s' if len(states) != 1 else ''} visited, "
            f"{len(entries)} trip{'s' if len(entries) != 1 else ''} logged."
        ).italic = True
    else:
        document.add_paragraph("No trips logged yet.")

    for entry in entries:
        document.add_heading(entry.get("title") or entry["state"], level=2)
        subtitle = document.add_paragraph()
        subtitle_bits = [entry["state"]]
        if entry.get("park_name"):
            subtitle_bits.append(entry["park_name"])
        subtitle_bits.append(entry["visited_on"])
        subtitle.add_run(" · ".join(subtitle_bits)).bold = True

        if entry.get("story"):
            document.add_paragraph(entry["story"])
        if entry.get("favorite_moment"):
            favorite = document.add_paragraph()
            favorite.add_run("Favorite moment: ").bold = True
            favorite.add_run(entry["favorite_moment"])
        if entry.get("would_return"):
            would_return = document.add_paragraph()
            would_return.add_run("Would go back? ").bold = True
            would_return.add_run(entry["would_return"])

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def course_filename(course: dict[str, Any]) -> str:
    """A readable .docx filename: the course title, slugged, plus today's date."""
    title = course.get("title") or "course"
    slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-") or "course"
    return f"{slug}-documentation-{date.today().isoformat()}.docx"


def course_to_docx(
    course: dict[str, Any],
    activities: list[dict[str, Any]],
    student_name: str,
) -> bytes:
    """Render one course's full district documentation packet.

    `activities` is `Database.course_activities(course_id)` -- each row is an
    `activities` record, carrying its full generated `lesson` (assignment
    content, assessment description, quiz result) when it came from one.
    Covers, in order, the seven pieces Sumner-Bonney Lake's grades 6-12
    packet requires: description, goals/objectives, outline, the hours log
    (150 hours = 1 credit), completed assignments/assessments, how
    performance is assessed, and progress + final grade (converted to
    Pass/Fail on the transcript).
    """
    document = Document()
    style = document.styles["Normal"]
    style.font.size = Pt(11)

    document.add_heading(course.get("title") or "Course", level=1)
    target_hours = round(config.CREDIT_HOURS_PER_UNIT * (course.get("credit_value") or 1.0), 1)
    subtitle = document.add_paragraph()
    subtitle.add_run(
        f"{student_name} · {subjects.label(course['credit_subject'])} · "
        f"{course.get('credit_value', 1.0):g} credit · Grade {course.get('grade_level') or '—'}"
    ).bold = True
    document.add_paragraph(f"{course['start_date']} through {course['end_date']}")

    document.add_heading("Course description", level=2)
    document.add_paragraph(course.get("description") or "—")

    document.add_heading("Course goals and objectives", level=2)
    document.add_paragraph(course.get("goals") or "—")

    document.add_heading("Course outline of the program", level=2)
    document.add_paragraph(course.get("outline") or "—")

    total_minutes = sum(a["minutes"] for a in activities)
    total_hours = round(total_minutes / 60, 1)
    document.add_heading("Learning activities and instructional time log", level=2)
    document.add_paragraph(
        f"{total_hours:g} of {target_hours:g} hours logged "
        f"({round(100 * total_hours / target_hours) if target_hours else 0}% of "
        f"{course.get('credit_value', 1.0):g} credit)."
    )
    if activities:
        table = document.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        header[0].text, header[1].text, header[2].text = "Date", "Activity", "Minutes"
        for activity in activities:
            row = table.add_row().cells
            row[0].text = activity["occurred_on"]
            row[1].text = activity["title"]
            row[2].text = str(activity["minutes"])
    else:
        document.add_paragraph("No instructional time logged toward this course yet.")

    document.add_heading("Completed assignments and assessments", level=2)
    if activities:
        for activity in activities:
            lesson = activity.get("lesson")
            heading = f"{activity['occurred_on']} — {activity['title']}"
            document.add_heading(heading, level=3)
            payload = lesson["payload"] if lesson else {}
            for objective in payload.get("learning_objectives") or []:
                document.add_paragraph(str(objective), style="List Bullet")
            if activity.get("description"):
                document.add_paragraph(activity["description"])
            quiz_result = (lesson or {}).get("metadata", {}).get("quiz_result")
            if quiz_result and quiz_result.get("total"):
                verdict = "passed" if quiz_result.get("passed") else "did not yet pass"
                document.add_paragraph(
                    f"Quiz: {quiz_result['correct']}/{quiz_result['total']} — {verdict}."
                )
    else:
        document.add_paragraph("Nothing completed toward this course yet.")

    document.add_heading("How student performance is assessed", level=2)
    described_any = False
    seen: set[tuple[str, str]] = set()
    for activity in activities:
        lesson = activity.get("lesson")
        assessment = (lesson or {}).get("payload", {}).get("assessment") or {}
        key = (assessment.get("kind", ""), assessment.get("description", ""))
        if not assessment or key in seen:
            continue
        seen.add(key)
        described_any = True
        if assessment.get("kind"):
            document.add_paragraph(assessment["kind"]).runs[0].bold = True
        if assessment.get("description"):
            document.add_paragraph(assessment["description"])
        if assessment.get("mastery_criteria"):
            criteria = document.add_paragraph()
            criteria.add_run("Mastery: ").bold = True
            criteria.add_run(assessment["mastery_criteria"])
    if not described_any:
        document.add_paragraph(
            "Performance assessed through direct parent observation of each completed "
            "activity, checked against the activity's stated purpose."
        )

    document.add_heading("Student progress and final grade", level=2)
    document.add_paragraph(f"{total_hours:g} of {target_hours:g} hours completed.")
    if course.get("final_grade"):
        document.add_paragraph(f"Final grade: {course['final_grade']}")
    if course.get("pass_fail"):
        document.add_paragraph(f"Transcript record: {course['pass_fail'].upper()}")
    else:
        document.add_paragraph("Transcript record: in progress — not yet finalized.")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class DocxExtractionError(Exception):
    """The uploaded file isn't a .docx python-docx can actually open."""


def extract_docx_text(file: Any) -> str:
    """Pull plain text back out of an uploaded .docx -- one paragraph per
    line, plus any tables (read row by row, cells joined with " | ") since a
    written response could reasonably land in either. `file` is anything
    python-docx's own `Document()` accepts: a path, bytes, or a file-like
    object (Streamlit's `st.file_uploader` result works directly).

    Raises `DocxExtractionError` on anything that isn't a real .docx, so the
    caller can show a plain "that didn't look like a Word doc" message
    instead of a raw traceback -- an uploaded .doc, .pdf, or a corrupted
    file all land here rather than crashing the page.
    """
    try:
        document = Document(file)
    except Exception as exc:
        raise DocxExtractionError(
            "That doesn't look like a valid Word (.docx) file."
        ) from exc
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines).strip()
